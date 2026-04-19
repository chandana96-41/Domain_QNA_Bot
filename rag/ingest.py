"""Load domain files, chunk, and add to Chroma."""
from pathlib import Path
from typing import List

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_core.vectorstores import VectorStoreRetriever
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from rag.config import CHROMA_COLLECTION, CHROMA_DIR, DEFAULT_EMBED_MODEL, ollama_base_url


def _load_one(path: Path) -> List[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        loader = PyPDFLoader(str(path))
    elif suffix in (".txt", ".md"):
        loader = TextLoader(str(path), encoding="utf-8")
    elif suffix == ".docx":
        from docx import Document as DocxDocument

        d = DocxDocument(str(path))
        text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata={"source_file": path.name})]
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Use .pdf, .txt, .md, or .docx")

    docs = loader.load()
    for d in docs:
        d.metadata["source_file"] = path.name
    return docs


def load_and_split_files(paths: List[Path]) -> List[Document]:
    all_docs: List[Document] = []
    for p in paths:
        all_docs.extend(_load_one(p))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(all_docs)


def get_embeddings(model: str | None = None) -> OllamaEmbeddings:
    return OllamaEmbeddings(
        model=model or DEFAULT_EMBED_MODEL,
        base_url=ollama_base_url(),
    )


def _get_store_retriever(embedding: OllamaEmbeddings) -> VectorStoreRetriever | None:
    """
    Return retriever if a valid persisted collection exists, else None.
    This avoids assuming any non-empty directory is a healthy Chroma collection.
    """
    if not CHROMA_DIR.exists() or not any(CHROMA_DIR.iterdir()):
        return None
    try:
        store = Chroma(
            collection_name=CHROMA_COLLECTION,
            embedding_function=embedding,
            persist_directory=str(CHROMA_DIR),
        )
        return store.as_retriever(search_kwargs={"k": 1})
    except Exception:
        return None


def ingest_files(paths: List[Path], embed_model: str | None = None, replace: bool = False) -> int:
    """
    Add documents from paths into Chroma. If replace=True, wipes existing index first.
    If index already exists and replace=False, appends new chunks.
    """
    if not paths:
        return 0

    if replace:
        clear_index()

    chunks = load_and_split_files(paths)
    if not chunks:
        return 0

    emb = get_embeddings(embed_model)
    existing = _get_store_retriever(emb)
    if existing is not None:
        store = Chroma(
            collection_name=CHROMA_COLLECTION,
            embedding_function=emb,
            persist_directory=str(CHROMA_DIR),
        )
        store.add_documents(chunks)
    else:
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        Chroma.from_documents(
            documents=chunks,
            embedding=emb,
            persist_directory=str(CHROMA_DIR),
            collection_name=CHROMA_COLLECTION,
        )
    return len(chunks)


def clear_index() -> None:
    """Remove persisted Chroma data (fresh start)."""
    import shutil

    if CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
